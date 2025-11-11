#!/usr/bin/env python3
"""
Inspect the structure of the Word document to understand how to parse it.
"""

from docx import Document

def inspect_document(docx_path):
    """Print the structure of the document."""
    doc = Document(docx_path)
    
    print("=" * 80)
    print("Document Structure Inspection")
    print("=" * 80)
    
    print(f"\nTotal paragraphs: {len(doc.paragraphs)}")
    print(f"Total tables: {len(doc.tables)}")
    
    print("\n" + "=" * 80)
    print("First 30 paragraphs:")
    print("=" * 80)
    
    for i, para in enumerate(doc.paragraphs[:30]):
        text = para.text.strip()
        if not text:
            continue
        
        # Check if bold
        is_bold = False
        if para.runs:
            is_bold = any(run.bold for run in para.runs)
        
        # Check style
        style = para.style.name if para.style else "None"
        
        print(f"\n[{i}] Style: {style}, Bold: {is_bold}")
        print(f"Text: {text[:100]}{'...' if len(text) > 100 else ''}")
    
    # Check if there are tables
    if doc.tables:
        print("\n" + "=" * 80)
        print("Tables found in document:")
        print("=" * 80)
        for i, table in enumerate(doc.tables[:3]):
            print(f"\nTable {i}: {len(table.rows)} rows x {len(table.columns)} columns")
            if table.rows:
                print("First few rows:")
                for row_idx, row in enumerate(table.rows[:5]):
                    cells_text = [cell.text.strip()[:50] for cell in row.cells]
                    print(f"  Row {row_idx}: {cells_text}")

if __name__ == "__main__":
    inspect_document("Aggregate Mini Wiki.docx")
